"""
Generates the Parking Digital Twin Implementation Review as a Word document.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Helpers ───────────────────────────────────────────────────────────────────
def shade_cell(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)

def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x1A, 0x37, 0x6C)
    return p

def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x1A, 0x37, 0x6C)
    return p

def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def add_body(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

def add_table(headers, rows, col_widths=None, header_bg='1A376C'):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        shade_cell(hdr_cells[i], header_bg)
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
    for ri, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        for ci, cell_text in enumerate(row_data):
            row_cells[ci].text = str(cell_text)
            if ri % 2 == 1:
                shade_cell(row_cells[ci], 'EBF3FB')
            for run in row_cells[ci].paragraphs[0].runs:
                run.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)
    doc.add_paragraph()
    return table


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("UNB Parking Digital Twin")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = RGBColor(0x1A, 0x37, 0x6C)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Implementation Review — Achievements, Architecture & Challenges")
r2.font.size = Pt(14)
r2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f"Generated: {datetime.date.today().strftime('%B %d, %Y')}").font.size = Pt(10)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 1. ARCHITECTURE OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
add_h1("1. Architecture Overview")
add_body(
    "The Parking Digital Twin is a full-stack TypeScript monorepo that models, "
    "simulates, and predicts parking demand for the UNB Saint John campus. It "
    "combines real academic-schedule data with a hybrid probabilistic prediction "
    "engine to support both real-time visualisation and what-if scenario planning."
)
add_h2("Technology Stack")
add_table(
    ["Layer", "Technology", "Purpose"],
    [
        ["Backend runtime",    "Node.js 20 + Express 4",                "REST API, middleware, simulator"],
        ["Language",           "TypeScript 5.6",                         "Type safety across BE & FE"],
        ["ORM / Database",     "TypeORM 0.3 + SQLite / PostgreSQL",     "Entities, queries, relations"],
        ["Frontend",           "React 18 + Vite 5 + TailwindCSS",       "SPA, maps, scenario UI"],
        ["Maps",               "Leaflet 1.9 + react-leaflet",            "Interactive campus map"],
        ["Caching",            "Redis 5 (optional)",                     "5-min / 2-min TTL caching"],
        ["Satellite imagery",  "Google Earth Engine",                    "Satellite tile overlay"],
        ["Auth",               "JWT + bcrypt",                           "Stateless auth, role-based access"],
        ["Testing",            "Jest (BE) + Vitest (FE)",                "Unit & integration tests"],
        ["API spec",           "OpenAPI 3.0.3",                          "Machine-readable REST contract"],
    ],
    col_widths=[1.8, 2.4, 3.0]
)


# ══════════════════════════════════════════════════════════════════════════════
# 2. WHAT HAS BEEN IMPLEMENTED
# ══════════════════════════════════════════════════════════════════════════════
add_h1("2. What Has Been Implemented")

# 2.1
add_h2("2.1  Core Parking System")
add_body(
    "Sixteen parking lots with approximately 1,231 individual spots are seeded from "
    "SVG files stored in the frontend. Each SVG encodes spot positions and labels; "
    "the seed script parses them to create ParkingSpot records with a 1:1 slotIndex "
    "mapping. Key spot attributes include distanceFromExit (used for recommendations) "
    "and isAccessible (disability flag). A ParkingSpotLog entity records every status "
    "change as a time-series for historical analysis and log browsing."
)
add_table(
    ["Entity", "Purpose"],
    [
        ["ParkingLot",     "Lot name, campus, total capacity, type (general / staff / resident / timed / phd)"],
        ["ParkingSpot",    "Individual stall — label, section, row, slotIndex, distanceFromExit, isAccessible, currentStatus"],
        ["ParkingSpotLog", "Time-series of occupied / empty status changes with timestamps"],
    ],
    col_widths=[2.0, 5.2]
)

# 2.2
add_h2("2.2  In-Process Parking Simulator")
add_body(
    "The simulator runs inside the backend process on a 5-second tick. Occupancy "
    "churn is calculated from a campus hourly-demand signal (campusOccupancyProfile.ts) "
    "with Gaussian noise. Quiet hours (22:00–07:00) reduce activity. A configurable "
    "floor (SIM_MIN_OCCUPIED_SPOTS) prevents lots from going completely empty."
)
add_body("Two map modes are supported:", bold=True)
add_bullet("Live — uses the Moncton real-time clock; the map reflects current simulated conditions.")
add_bullet("Scenario — a fixed date/time is locked in, enabling deterministic what-if visualisation.")
add_body(
    "The scenario clock is the key mechanism behind the day-long parking plan: clicking "
    "a plan segment sets the map to that timestamp so the user sees predicted parking "
    "conditions at that exact moment."
)

# 2.3
add_h2("2.3  Hybrid Prediction Engine")
add_body(
    "The most technically sophisticated component. A five-layer stacked model computes "
    "occupancy percentage for any lot at any datetime. Results are cached for 5 minutes."
)
add_table(
    ["Layer", "Source", "Logic"],
    [
        ["1 — Historical data",  "historical_proxy_data table",    "If ≥ 3 samples for (lot, hour, day-of-week, period) → use average"],
        ["2 — Hardcoded curves", "prediction.service.ts",          "24-hour occupancy curves per lot type (general, staff, resident, timed, phd)"],
        ["3 — DDM correction",   "lot_occupancy_correction table", "Residual = observed − predicted; weighted by tanh(nSamples / 10)"],
        ["4 — Activity curve",   "activityCurve.service.ts",       "Enrollment-driven demand index; lots near course buildings get a higher multiplier"],
        ["5 — Event boost",      "Query param eventSize",          "small +10 %, medium +20 %, large +35 % of remaining free capacity"],
    ],
    col_widths=[1.8, 2.2, 3.2]
)
add_body(
    "The academic calendar (hardcoded 2024–2026) classifies each date as one of: "
    "classes, reading_week, exams, holiday, pre_semester, or summer — each with its "
    "own multiplier. The final result is clamped to [0, 100] %."
)
add_body("Endpoints:", bold=True)
add_bullet("GET /api/prediction/lots/:lotId — single lot at a moment")
add_bullet("GET /api/prediction/lots/:lotId/day-profile — 24-hour occupancy curve")
add_bullet("GET /api/prediction/lots/:lotId/next-hours — N hours ahead")
add_bullet("GET /api/prediction/snapshot — all lots at a single moment")

# 2.4
add_h2("2.4  Campus Parameters (Behavioural Model)")
add_body(
    "A key-value store auto-seeded on startup holding the behavioural constants derived "
    "from the staff and student presence assumption documents."
)
add_table(
    ["Parameter", "Value", "Meaning"],
    [
        ["carpool_rate",          "0.12", "12 % of drivers share a car"],
        ["non_driver_rate",       "0.35", "35 % walk, cycle, or take the bus"],
        ["effective_driver_rate", "0.53", "Computed: 1 − non_driver − carpool / 2"],
        ["absence_rate",          "0.15", "15 % daily absenteeism"],
        ["friday_absence_mult",   "1.33", "33 % higher absenteeism on Fridays"],
        ["monday_absence_mult",   "1.13", "13 % higher absenteeism on Mondays"],
    ],
    col_widths=[2.4, 1.2, 3.6]
)
add_body(
    "getDemandMultiplier(dayOfWeek) combines these to scale the activity curve used by "
    "the prediction engine. There is currently no admin UI to edit these parameters."
)

# 2.5
add_h2("2.5  What-If Explorer")
add_body(
    "Calls predictSnapshot() twice — baseline (eventSize=none, useEnrollment=false) and "
    "scenario (user-selected parameters) — then computes per-lot deltas. Results are "
    "cached for 2 minutes."
)
add_body("Frontend views:", bold=True)
add_bullet("Campus view — all lots in a comparison table showing occupancy % and free-spot delta.")
add_bullet("My Scenario — personal card (auth required) showing when an event forces a different lot or triggers a 'fewer than 10 spots' warning.")

# 2.6
add_h2("2.6  Day-Long Arrival Planning")
add_body(
    "Given a student's courses for a chosen date, the arrival recommendation service "
    "builds named plan segments and resolves a parking recommendation for each one."
)
add_table(
    ["Segment type", "Description"],
    [
        ["initial_arrival",  "First arrival before the first class of the day"],
        ["stay_on_campus",   "Short gaps between classes — student remains on campus"],
        ["return_and_park",  "Gaps long enough to leave campus and return"],
    ],
    col_widths=[2.2, 5.0]
)
add_body(
    "For each segment the service: (1) queries the prediction API for expected occupancy, "
    "(2) runs the recommendation engine to find the best available lot and spot, "
    "(3) estimates walk time to the destination building including floor navigation, and "
    "(4) records an apply-scenario timestamp. Clicking a segment in the UI sets the map "
    "to that timestamp."
)

# 2.7
add_h2("2.7  Authentication & Role-Based Access")
add_body(
    "Stateless JWT authentication. On login the backend issues a signed token stored in "
    "browser localStorage. Roles (staff, student, phd_candidate) and flags (resident, "
    "disabled) are embedded in the token and used by the lot eligibility logic to filter "
    "which lots a user may park in."
)

# 2.8
add_h2("2.8  Historical Data Pipeline")
add_table(
    ["Script", "Purpose"],
    [
        ["generateHistoricalData.ts",  "Produces synthetic occupancy rows in historical_proxy_data"],
        ["generateResiduals.ts",        "Computes DDM corrections per (lot, hour, day-of-week, period)"],
        ["generateEventResiduals.ts",   "Event-specific residual corrections"],
        ["populateSpotAttributes.ts",   "Sets distanceFromExit and isAccessible per spot"],
        ["recalculateSpotDistances.ts", "Recalculates lot-to-building walking distances"],
        ["importBirminghamData.ts",     "Imports external Birmingham pilot dataset"],
        ["scrape-courses.ts",           "Scrapes UNB self-service course catalogue (requires auth token)"],
    ],
    col_widths=[3.0, 4.2]
)

# 2.9
add_h2("2.9  Google Earth Engine Integration")
add_body(
    "The backend proxies three Earth Engine endpoints (thumbnail, tiles, mapid) so the "
    "frontend can render satellite imagery without exposing service-account credentials. "
    "Initialisation is skipped gracefully if credentials are absent, falling back to "
    "OpenStreetMap tiles."
)

# 2.10
add_h2("2.10  Test Coverage")
add_table(
    ["Test file", "Framework", "What is covered"],
    [
        ["arrivalRecommendation.utils.test.ts", "Jest",   "Date parsing, floor inference, day-of-week helpers"],
        ["courseMeetingTime.test.ts",            "Jest",   "Course start/end time validation"],
        ["parkingLotEligibility.test.ts",        "Jest",   "Role-based lot access logic"],
        ["whatif.test.tsx",                       "Vitest", "What-If page rendering and API interaction"],
        ["apiClient.test.ts",                    "Vitest", "HTTP client error handling"],
    ],
    col_widths=[3.2, 1.3, 2.7]
)
add_body(
    "Estimated coverage: ~30 % of the codebase. The prediction engine, activity curve "
    "service, campus parameters, simulator, and arrival recommendation core logic "
    "are not covered by automated tests."
)


# ══════════════════════════════════════════════════════════════════════════════
# 3. SHORTCOMINGS AND CHALLENGES
# ══════════════════════════════════════════════════════════════════════════════
add_h1("3. Shortcomings and Challenges")

add_h2("3.1  Critical Data Issues")

add_h3("All historical data is synthetic")
add_body(
    "The prediction engine's 'data' confidence path (≥ 3 historical samples → use average) "
    "is trained entirely on fabricated data generated from the same occupancy curves it is "
    "supposed to improve on. Layers 1 and 2 of the prediction stack are correlated: if the "
    "base curves are wrong, the synthetic historical data is wrong in the same direction, and "
    "the DDM residual correction has nothing real to correct against. Predictions will look "
    "plausible but have no grounding in actual sensor observations."
)

add_h3("No real-time sensor integration")
add_body(
    "There are no parking sensors, cameras, or any external data feed. Every ParkingSpotLog "
    "entry is created by the simulator. The 'live' map is a probabilistic simulation, not a "
    "live feed. This is acceptable for a demonstration but is the largest gap for operational "
    "deployment."
)

add_h3("Course scraping is fragile")
add_body(
    "The scraper depends on the HTML structure of the UNB self-service portal and requires an "
    "active session token and cookie supplied manually. Any portal redesign breaks it silently. "
    "Scraped enrollment numbers become stale as the term progresses."
)

add_h2("3.2  Architectural Weaknesses")

add_h3("Academic calendar hardcoded and expires after 2026")
add_body(
    "Period date ranges (classes, reading week, exams, holiday, pre-semester, summer) are "
    "hardcoded in campusOccupancyProfile.ts through 2026. There is no admin interface or "
    "config file to extend the calendar. After 2026 every date falls into the default period "
    "silently."
)

add_h3("demandParameters.ts was never created")
add_body(
    "The implementation plan (Step 0) called for a single BE/src/config/demandParameters.ts "
    "to centralise all constants from the assumption documents. This file does not exist. "
    "Carpool rates, absence rates, and effective driver rates are hardcoded defaults in the "
    "campusParameter.service (under prediction) with no single place to audit all downstream effects."
)

add_h3("No database migrations — synchronize: true in production")
add_body(
    "The TypeORM data source uses synchronize: true, which auto-modifies the schema whenever "
    "entity definitions change. This is convenient during development but dangerous in "
    "production: it can silently drop columns or alter indexes. There are no migration files."
)

add_h3("Simulator is non-deterministic for a given scenario")
add_body(
    "Even in scenario mode (fixed date/time), the simulator continues its 5-second churn loop. "
    "Two users viewing the same scenario at different wall-clock times see different spot "
    "configurations because the simulator has continued to run between their visits."
)

add_h3("What-If module has no service layer")
add_body(
    "Unlike every other module — which follows entity / service / controller / route — the "
    "what-if module places all business logic directly in the route handler (whatif.route.ts). "
    "This makes unit testing impossible without spinning up the full Express stack and prevents "
    "the logic from being reused by other modules."
)

add_h3("Campus Parameters have no HTTP routes")
add_body(
    "Under modules/prediction, campusParameter.entity.ts and campusParameter.service.ts exist "
    "but no controller or route is registered. Behavioural parameters can only be changed by "
    "direct database access or by reseeding — no admin user can adjust them at runtime without "
    "a code deployment."
)

add_h3("Prediction confidence is binary, not probabilistic")
add_body(
    "PredictionResult.confidence is either 'data' or 'curve'. A lot with exactly 3 historical "
    "samples receives the same 'data' confidence label as one with 500 samples. Users and "
    "downstream consumers have no signal about actual prediction reliability."
)

add_h2("3.3  Testing Gaps")

add_h3("Prediction engine is completely untested")
add_body(
    "The five-layer hybrid model, DDM residual weighting (tanh formula), activity curve "
    "computation, and academic calendar lookups have zero automated tests. A change to any "
    "of these components would go undetected until a human noticed incorrect predictions."
)

add_h3("Simulator and arrival recommendation core logic are untested")
add_body(
    "The churn model, quiet-hours enforcement, and occupancy floor logic in the simulator "
    "are not tested. The arrivalRecommendation service test file covers utility helpers "
    "but not the plan-building algorithm itself."
)

add_h2("3.4  Operational Risks")

add_h3("SQLite in development vs PostgreSQL in production")
add_body(
    "The two databases behave differently in edge cases (type coercion, LIKE case-sensitivity, "
    "JSON column handling, concurrency). All automated tests run against SQLite or no database "
    "at all. A query that works locally may fail silently in production."
)

add_h3("Redis failure is silent under load")
add_body(
    "If Redis is unavailable the caching middleware silently no-ops. Under load each what-if "
    "request triggers two full predictSnapshot() calls, each querying historical data, "
    "correction tables, campus parameters, and course enrollment. There is no circuit breaker "
    "or rate limiter to protect the database."
)

add_h3("Earth Engine credentials have no expiry alerting")
add_body(
    "If the service account credentials expire or are revoked, the backend logs a warning at "
    "startup but continues running. Satellite imagery silently disappears with no UI indicator, "
    "no fallback message, and no monitoring hook."
)


# ══════════════════════════════════════════════════════════════════════════════
# 4. SUMMARY SCORECARD
# ══════════════════════════════════════════════════════════════════════════════
add_h1("4. Summary Scorecard")
add_table(
    ["Area", "Status", "Biggest Risk / Gap"],
    [
        ["Core parking CRUD",        "Complete",          "None"],
        ["In-process simulator",     "Complete",          "Non-deterministic in scenario mode"],
        ["Hybrid prediction engine", "Complete",          "Trained on synthetic data only"],
        ["What-If explorer",         "Complete",          "No service layer; untestable in isolation"],
        ["Day-long arrival plan",    "Complete",          "Scenario clock drifts between clicks"],
        ["Campus parameters",        "Backend only",      "No HTTP routes, no admin UI"],
        ["Historical data",          "Synthetic only",    "No real sensor observations"],
        ["Academic calendar",        "Hardcoded",         "Expires silently after 2026"],
        ["Database migrations",      "Not implemented",   "synchronize: true dangerous in production"],
        ["Test coverage",            "~30 %",             "Critical business logic untested"],
        ["Real-time sensor data",    "Not implemented",   "All 'live' data is simulated"],
        ["Campus parameters UI",     "Not implemented",   "Requires direct DB access to change"],
    ],
    col_widths=[2.4, 1.8, 3.0]
)


# ══════════════════════════════════════════════════════════════════════════════
# 5. RECOMMENDED NEXT STEPS
# ══════════════════════════════════════════════════════════════════════════════
add_h1("5. Recommended Next Steps")

add_h2("High Priority")
add_bullet("Add TypeORM migration files and disable synchronize: true for production.")
add_bullet("Write unit tests for the prediction engine (all five layers), activity curve service, and arrival recommendation core algorithm — target >= 70 % coverage.")
add_bullet("Extract what-if business logic into a WhatIfService class to enable unit testing and reuse.")
add_bullet("Expose CampusParameters via authenticated admin HTTP routes (GET / PATCH /api/campus-parameters).")

add_h2("Medium Priority")
add_bullet("Replace the hardcoded academic calendar with a database-backed table editable via an admin endpoint.")
add_bullet("Create BE/src/config/demandParameters.ts as the single source of truth for all behavioural constants.")
add_bullet("Add a numeric confidence score (sample count, standard deviation) to PredictionResult alongside the binary flag.")
add_bullet("Implement a Redis health-check warning and set a request rate limit on prediction and what-if endpoints.")

add_h2("Lower Priority / Future Work")
add_bullet("Integrate real occupancy data (sensors, camera inference, or manual counts) to replace synthetic historical data.")
add_bullet("Add snapshot isolation to the simulator in scenario mode so the spot distribution is frozen when a scenario is set.")
add_bullet("Replace the fragile UNB self-service scraper with a direct Banner API integration or a manually-maintained CSV import.")
add_bullet("Add satellite imagery health monitoring so the UI displays a fallback indicator when Earth Engine credentials fail.")


# ══════════════════════════════════════════════════════════════════════════════
# Save
# ══════════════════════════════════════════════════════════════════════════════
output_path = r"c:\Users\obivi\Parking_Digital_Twin\docs\ParkingDigitalTwin_ImplementationReview.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
